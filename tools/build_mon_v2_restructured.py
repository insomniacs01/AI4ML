from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


BASE = Path(r"D:\333\AI4ML")
SRC = BASE / "MON-\u667a\u7b97\u793e\u533a-2-\u9700\u6c42-0413.docx"
OUT_NAMES = [
    "MON-\u667a\u7b97\u793e\u533a-2-\u9700\u6c42-V2-\u91cd\u6784\u7248.docx",
    "MON-\u667a\u7b97\u793e\u533a-2-\u9700\u6c42-V2-\u4fdd\u7559\u539f\u56fe\u91cd\u6784\u7248.docx",
]


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Paragraph not found: {text}")


def replace_paragraph(doc: Document, old: str, new: str, style: str | None = None) -> Paragraph:
    para = find_paragraph(doc, old)
    para.text = new
    if style:
        para.style = style
    return para


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_bullets_after(paragraph: Paragraph, bullets: list[str]) -> None:
    last = paragraph
    for bullet in bullets:
        last = insert_paragraph_after(last, bullet, "List Bullet")


def replace_all(doc: Document, old: str, new: str) -> None:
    for para in doc.paragraphs:
        if old in para.text:
            para.text = para.text.replace(old, new)


def build() -> None:
    doc = Document(str(SRC))

    # Cover/version
    replace_paragraph(doc, "\u9700\u6c42\u89c4\u683c\u8bf4\u660e\u4e66", "\u9700\u6c42\u89c4\u683c\u8bf4\u660e\u4e66 V2")
    replace_paragraph(doc, "\u4ea7\u54c1/\u8f6f\u4ef6\u9700\u6c42\u89c4\u683c\u4e66", "\u4ea7\u54c1/\u8f6f\u4ef6\u9700\u6c42\u89c4\u683c\u4e66")
    info = doc.tables[0]
    info.cell(0, 2).text = "AI4ML-FR-2"
    info.cell(1, 2).text = "V2"
    info.cell(3, 2).text = "2026-04-16"

    # Product expression
    replace_paragraph(
        doc,
        "AI4ML \u662f\u4e00\u4e2a\u5206\u9636\u6bb5\u6f14\u8fdb\u7684\u5e73\u53f0\uff1a",
        "AI4ML \u662f\u4e00\u4e2a\u9762\u5411\u6559\u5b66\u3001\u5b9e\u9a8c\u4e0e\u539f\u578b\u9a8c\u8bc1\u573a\u666f\u7684 AI \u673a\u5668\u5b66\u4e60\u5de5\u4f5c\u53f0\uff0c\u5e76\u9010\u6b65\u6f14\u8fdb\u4e3a\u652f\u6301\u56e2\u961f\u6cbb\u7406\u3001\u8fc7\u7a0b\u534f\u540c\u548c\u5de5\u4f5c\u6d41\u590d\u7528\u7684\u5e73\u53f0\uff1a",
    )

    # Move architecture chapter into chapter 2 as a subsection
    replace_paragraph(
        doc,
        "3 \u603b\u4f53\u67b6\u6784\u4e0e\u4e1a\u52a1\u6d41\u7a0b",
        "2.8 \u603b\u4f53\u67b6\u6784\u4e0e\u4e1a\u52a1\u6d41\u7a0b",
        "Heading 2",
    )
    replace_paragraph(doc, "3.1 \u603b\u4f53\u7cfb\u7edf\u67b6\u6784\u56fe", "2.8.1 \u603b\u4f53\u7cfb\u7edf\u67b6\u6784\u56fe", "Heading 3")
    replace_paragraph(doc, "3.2 \u6838\u5fc3\u4e1a\u52a1\u4e3b\u6d41\u7a0b\u56fe", "2.8.2 \u6838\u5fc3\u4e1a\u52a1\u4e3b\u6d41\u7a0b\u56fe", "Heading 3")
    replace_paragraph(doc, "3.3 AI \u5de5\u4f5c\u6d41\u4e0e\u9636\u6bb5\u72b6\u6001\u56fe", "2.8.3 AI \u5de5\u4f5c\u6d41\u4e0e\u9636\u6bb5\u72b6\u6001\u56fe", "Heading 3")
    replace_paragraph(doc, "3.4 \u4eba\u673a\u534f\u540c\u5e72\u9884\u6d41\u7a0b\u56fe", "2.8.4 \u4eba\u673a\u534f\u540c\u5e72\u9884\u6d41\u7a0b\u56fe", "Heading 3")
    replace_paragraph(doc, "3.5 \u7ba1\u7406\u5458\u6cbb\u7406\u6d41\u7a0b\u56fe", "2.8.5 \u7ba1\u7406\u5458\u6cbb\u7406\u6d41\u7a0b\u56fe", "Heading 3")
    replace_paragraph(doc, "3.6 \u4efb\u52a1\u4e0e\u5de5\u4f5c\u6d41\u72b6\u6001\u6d41\u8f6c", "2.8.6 \u4efb\u52a1\u4e0e\u5de5\u4f5c\u6d41\u72b6\u6001\u6d41\u8f6c", "Heading 3")
    for old, new in [
        ("\u56fe 3-1", "\u56fe 2-1"),
        ("\u56fe 3-2", "\u56fe 2-2"),
        ("\u56fe 3-3", "\u56fe 2-3"),
        ("\u56fe 3-4", "\u56fe 2-4"),
        ("\u56fe 3-5", "\u56fe 2-5"),
    ]:
        replace_all(doc, old, new)

    # Functional chapter restructure
    replace_paragraph(doc, "4 \u529f\u80fd\u6027\u9700\u6c42", "3 \u529f\u80fd\u6027\u9700\u6c42", "Heading 1")
    p = replace_paragraph(doc, "4.1 \u4f18\u5148\u7ea7\u5b9a\u4e49", "3.1 \u529f\u80fd\u6027\u9700\u6c42\u5206\u7c7b\u4e0e\u4f18\u5148\u7ea7", "Heading 2")
    insert_bullets_after(
        p,
        [
            "\u57fa\u7840\u652f\u6491\u4e0e\u4efb\u52a1\u5e95\u5ea7\uff1a\u8986\u76d6\u7cfb\u7edf\u72b6\u6001\u3001\u4efb\u52a1\u5bf9\u8c61\u548c\u672c\u5730\u6267\u884c\u94fe\u8def\u3002",
            "\u56e2\u961f\u6cbb\u7406\u4e0e\u8d44\u6e90\u7ba1\u7406\uff1a\u8986\u76d6\u6210\u5458\u3001\u6743\u9650\u3001\u989d\u5ea6\u3001AI \u8fde\u63a5\u5668\u4e0e\u8d44\u4ea7\u7ba1\u7406\u3002",
            "AI \u4efb\u52a1\u5de5\u4f5c\u53f0\u4e0e\u7ed3\u679c\u5206\u6790\uff1a\u8986\u76d6\u4efb\u52a1\u53d1\u8d77\u3001\u6570\u636e\u4e0a\u4f20\u3001\u5de5\u4f5c\u6d41\u6267\u884c\u3001\u8fdb\u5ea6\u5c55\u793a\u4e0e\u62a5\u544a\u8f93\u51fa\u3002",
            "\u4eba\u673a\u534f\u540c\u4e0e\u5de5\u4f5c\u6d41\u8d44\u4ea7\uff1a\u8986\u76d6\u4eba\u5de5\u5e72\u9884\u3001\u4ee3\u7801\u900f\u660e\u3001\u5206\u4eab\u53d1\u5e03\u4e0e Fork \u590d\u7528\u3002",
        ],
    )
    replace_paragraph(doc, "4.2 \u5e73\u53f0\u57fa\u7840\u652f\u6491\u6a21\u5757", "3.2 \u57fa\u7840\u652f\u6491\u4e0e\u4efb\u52a1\u5e95\u5ea7\u6a21\u5757", "Heading 2")
    replace_paragraph(doc, "4.3 \u7ba1\u7406\u5458\u7aef\u529f\u80fd\u9700\u6c42", "3.3 \u56e2\u961f\u6cbb\u7406\u4e0e\u8d44\u6e90\u7ba1\u7406\u6a21\u5757", "Heading 2")
    replace_paragraph(doc, "4.4 \u96f6\u57fa\u7840\u4e1a\u52a1\u7528\u6237\u529f\u80fd\u9700\u6c42", "3.4 AI \u4efb\u52a1\u5de5\u4f5c\u53f0\u4e0e\u7ed3\u679c\u5206\u6790\u6a21\u5757", "Heading 2")
    replace_paragraph(doc, "4.5 \u4eba\u673a\u534f\u540c\u4e0e\u5f00\u53d1\u589e\u5f3a\u9700\u6c42", "3.5 \u4eba\u673a\u534f\u540c\u4e0e\u5de5\u4f5c\u6d41\u8d44\u4ea7\u6a21\u5757", "Heading 2")

    # Data chapter renumber
    replace_paragraph(doc, "5 \u6570\u636e\u9700\u6c42", "4 \u6570\u636e\u9700\u6c42", "Heading 1")
    replace_paragraph(doc, "5.1 \u6838\u5fc3\u6570\u636e\u5bf9\u8c61", "4.1 \u6838\u5fc3\u6570\u636e\u5bf9\u8c61", "Heading 2")
    replace_paragraph(doc, "5.2 TaskRecord \u5b57\u6bb5\u8981\u6c42", "4.2 TaskRecord \u5b57\u6bb5\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "5.3.7 HumanInteractionRequest", "4.3 HumanInteractionRequest \u5b57\u6bb5\u8981\u6c42", "Heading 2")
    replace_paragraph(
        doc,
        "5.3 TokenLedger \u4e0e QuotaAccount \u5b57\u6bb5\u8981\u6c42",
        "4.4 TokenLedger\u3001QuotaAccount \u4e0e\u7ec4\u7ec7\u6cbb\u7406\u5bf9\u8c61\u5b57\u6bb5\u8981\u6c42",
        "Heading 2",
    )
    replace_paragraph(doc, "5.3.1 QuotaAccount", "4.4.1 QuotaAccount", "Heading 3")
    replace_paragraph(doc, "5.3.2 TokenLedger", "4.4.2 TokenLedger", "Heading 3")
    replace_paragraph(doc, "5.3.3 Team", "4.4.3 Team", "Heading 3")
    replace_paragraph(doc, "5.3.4 TeamMember", "4.4.4 TeamMember", "Heading 3")
    replace_paragraph(doc, "5.3.5 AIConnector", "4.4.5 AIConnector", "Heading 3")
    replace_paragraph(doc, "5.3.6 AIRoutingPolicy", "4.4.6 AIRoutingPolicy", "Heading 3")
    replace_paragraph(doc, "5.4 \u8d44\u4ea7\u6570\u636e\u5bf9\u8c61\u8981\u6c42", "4.5 \u8d44\u4ea7\u6570\u636e\u5bf9\u8c61\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "5.5 \u6570\u636e\u89c4\u5219\u4e0e\u7ea6\u675f", "4.6 \u6570\u636e\u89c4\u5219\u4e0e\u7ea6\u675f", "Heading 2")
    replace_paragraph(doc, "5.6 \u5b58\u50a8\u7ed3\u6784\u8981\u6c42", "4.7 \u5b58\u50a8\u7ed3\u6784\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "5.6.1 \u6587\u4ef6\u4ea7\u7269\u5b58\u50a8\u7ed3\u6784", "4.7.1 \u6587\u4ef6\u4ea7\u7269\u5b58\u50a8\u7ed3\u6784", "Heading 3")
    replace_paragraph(doc, "5.6.2 \u5143\u6570\u636e\u4e0e\u805a\u5408\u8bbf\u95ee\u7ed3\u6784", "4.7.2 \u5143\u6570\u636e\u4e0e\u805a\u5408\u8bbf\u95ee\u7ed3\u6784", "Heading 3")

    # Interface/page chapter renumber
    replace_paragraph(doc, "6 \u63a5\u53e3\u4e0e\u9875\u9762\u9700\u6c42", "5 \u9875\u9762\u4e0e\u63a5\u53e3\u9700\u6c42", "Heading 1")
    replace_paragraph(doc, "6.1 \u524d\u7aef\u9875\u9762\u9700\u6c42", "5.1 \u9875\u9762\u9700\u6c42", "Heading 2")
    replace_paragraph(doc, "6.2 \u6838\u5fc3\u4efb\u52a1 API", "5.2 \u6838\u5fc3\u4efb\u52a1\u63a5\u53e3", "Heading 2")
    replace_paragraph(doc, "6.3 \u56e2\u961f\u6cbb\u7406\u4e0e\u534f\u540c API", "5.3 \u56e2\u961f\u6cbb\u7406\u4e0e\u534f\u540c\u63a5\u53e3", "Heading 2")
    replace_paragraph(doc, "6.4 \u63a5\u53e3\u8bbe\u8ba1\u539f\u5219", "5.4 \u63a5\u53e3\u8bbe\u8ba1\u539f\u5219", "Heading 2")

    # Non-functional renumber
    replace_paragraph(doc, "7 \u975e\u529f\u80fd\u6027\u9700\u6c42", "6 \u975e\u529f\u80fd\u6027\u9700\u6c42", "Heading 1")
    replace_paragraph(doc, "7.1 \u6613\u7528\u6027\u8981\u6c42", "6.1 \u6613\u7528\u6027\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "7.2 \u6027\u80fd\u8981\u6c42", "6.2 \u6027\u80fd\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "7.3 \u53ef\u9760\u6027\u8981\u6c42", "6.3 \u53ef\u9760\u6027\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "7.4 \u5b89\u5168\u6027\u8981\u6c42", "6.4 \u5b89\u5168\u6027\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "7.5 \u53ef\u7ef4\u62a4\u6027\u8981\u6c42", "6.5 \u53ef\u7ef4\u62a4\u6027\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "7.6 \u53ef\u6269\u5c55\u6027\u8981\u6c42", "6.6 \u53ef\u6269\u5c55\u6027\u8981\u6c42", "Heading 2")
    replace_paragraph(doc, "7.7 \u53ef\u5ba1\u8ba1\u6027\u8981\u6c42", "6.7 \u53ef\u5ba1\u8ba1\u6027\u8981\u6c42", "Heading 2")

    # Technical constraints / risk / acceptance renumber
    replace_paragraph(doc, "8 \u6280\u672f\u7ea6\u675f\u3001\u98ce\u9669\u4e0e\u9a8c\u6536", "7 \u6280\u672f\u7ea6\u675f\u3001\u98ce\u9669\u4e0e\u9a8c\u6536", "Heading 1")
    replace_paragraph(doc, "8.1 \u6280\u672f\u7ea6\u675f", "7.1 \u6280\u672f\u7ea6\u675f", "Heading 2")
    replace_paragraph(doc, "8.2 \u98ce\u9669\u5206\u6790", "7.2 \u98ce\u9669\u5206\u6790", "Heading 2")
    replace_paragraph(doc, "8.3 \u9a8c\u6536\u53e3\u5f84", "7.3 \u9a8c\u6536\u53e3\u5f84", "Heading 2")
    replace_paragraph(doc, "8.3.1 \u529f\u80fd\u4e0e\u6d41\u7a0b\u9a8c\u6536", "7.3.1 \u529f\u80fd\u4e0e\u6d41\u7a0b\u9a8c\u6536", "Heading 3")
    replace_paragraph(doc, "8.3.2 \u6cbb\u7406\u4e0e\u534f\u540c\u9a8c\u6536", "7.3.2 \u6cbb\u7406\u4e0e\u534f\u540c\u9a8c\u6536", "Heading 3")

    for out_name in OUT_NAMES:
        out = BASE / out_name
        if out.exists():
            out.unlink()
        doc.save(str(out))


if __name__ == "__main__":
    build()
    for name in OUT_NAMES:
        print(BASE / name)
