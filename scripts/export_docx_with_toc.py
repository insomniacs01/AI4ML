from __future__ import annotations

import argparse
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER


TABLE_SEPARATOR_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+\s*$")
HEADING_RE = re.compile(r"^(#+)\s+(.*)$")
BULLET_RE = re.compile(r"^-\s+(.*)$")
NUMBER_RE = re.compile(r"^\d+\.\s+(.*)$")


def configure_document(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Songti SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")

    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_inline(paragraph, text: str) -> None:
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j != -1:
                run = paragraph.add_run(text[i + 2 : j])
                run.bold = True
                i = j + 2
                continue

        if text.startswith("`", i):
            j = text.find("`", i + 1)
            if j != -1:
                run = paragraph.add_run(text[i + 1 : j])
                run.font.name = "Menlo"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
                i = j + 1
                continue

        next_pos = len(text)
        for token in ("**", "`"):
            pos = text.find(token, i)
            if pos != -1:
                next_pos = min(next_pos, pos)
        paragraph.add_run(text[i:next_pos])
        i = next_pos


def add_toc_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("目 录")
    run.bold = True
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    run.font.size = Pt(16)


def make_bookmark_name(index: int) -> str:
    return f"toc_anchor_{index}"


def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F5AA6")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), "Songti SC")
    r_fonts.set(qn("w:ascii"), "Songti SC")
    r_fonts.set(qn("w:hAnsi"), "Songti SC")
    r_pr.append(r_fonts)

    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def collect_headings(markdown: str) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        stripped = line.strip()
        match = HEADING_RE.match(stripped)
        if not match:
            continue
        level = len(match.group(1))
        text = match.group(2).strip()
        if level == 1:
            continue
        mapped_level = max(1, min(3, level - 1))
        headings.append((mapped_level, text))
    return headings


def estimate_heading_pages(markdown: str, start_page: int = 2, lines_per_page: float = 40.0) -> list[int]:
    lines = markdown.splitlines()
    page = start_page
    used = 0.0
    in_code = False
    i = 0
    heading_pages: list[int] = []

    def block_lines_for_text(text: str, chars_per_line: int = 34, extra: float = 0.5) -> float:
        if not text.strip():
            return 0.0
        visual_len = 0
        for ch in text:
            visual_len += 1 if ord(ch) > 127 else 1
        return max(1.0, math.ceil(visual_len / chars_per_line) + extra)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            block_height = 1.0
            if used + block_height > lines_per_page:
                page += 1
                used = 0.0
            used += block_height
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1].strip()):
            row_count = 1
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_count += 1
                i += 1
            block_height = row_count * 1.35 + 0.5
            while used + block_height > lines_per_page:
                page += 1
                used = 0.0
            used += block_height
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            if level == 1:
                i += 1
                continue
            text = heading_match.group(2).strip()
            block_height = block_lines_for_text(text, chars_per_line=26, extra=0.8)
            if used + block_height > lines_per_page:
                page += 1
                used = 0.0
            heading_pages.append(page)
            used += block_height
            i += 1
            continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            block_height = block_lines_for_text(bullet_match.group(1), chars_per_line=32, extra=0.4)
            if used + block_height > lines_per_page:
                page += 1
                used = 0.0
            used += block_height
            i += 1
            continue

        number_match = NUMBER_RE.match(stripped)
        if number_match:
            block_height = block_lines_for_text(number_match.group(1), chars_per_line=32, extra=0.4)
            if used + block_height > lines_per_page:
                page += 1
                used = 0.0
            used += block_height
            i += 1
            continue

        block_height = block_lines_for_text(stripped, chars_per_line=34, extra=0.5)
        if used + block_height > lines_per_page:
            page += 1
            used = 0.0
        used += block_height
        i += 1

    return heading_pages


def add_visible_toc(doc: Document, headings: list[tuple[int, str]], heading_pages: list[int]) -> None:
    add_toc_title(doc)
    for index, (level, text) in enumerate(headings, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.22 * (level - 1))
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.1),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.SPACES,
        )
        add_internal_hyperlink(paragraph, text, make_bookmark_name(index))
        paragraph.add_run("\t")
        page_run = paragraph.add_run(str(heading_pages[index - 1]))
        page_run.font.name = "Songti SC"
        page_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = table.cell(r_idx, c_idx).paragraphs[0]
            add_inline(paragraph, text)
            for run in paragraph.runs:
                run.font.size = Pt(10.5)
                if r_idx == 0:
                    run.bold = True


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    headings = collect_headings(markdown)
    heading_pages = estimate_heading_pages(markdown)
    in_code = False
    code_buffer: list[str] = []
    title_written = False
    toc_inserted = False
    heading_index = 0
    bookmark_id = 1
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                for idx, code_line in enumerate(code_buffer):
                    if idx:
                        paragraph.add_run("\n")
                    run = paragraph.add_run(code_line)
                    run.font.name = "Menlo"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
                    run.font.size = Pt(10)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1].strip()):
            table_lines = [line]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = [
                [cell.strip() for cell in row_line.strip().strip("|").split("|")]
                for row_line in table_lines
            ]
            add_table(doc, rows)
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()

            if not title_written and level == 1:
                paragraph = doc.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, heading_text)
                title_written = True

                if not toc_inserted:
                    doc.add_paragraph()
                    add_visible_toc(doc, headings, heading_pages)
                    doc.add_page_break()
                    toc_inserted = True

                i += 1
                continue

            mapped_level = max(1, min(3, level - 1))
            paragraph = doc.add_paragraph(style=f"Heading {mapped_level}")
            add_inline(paragraph, heading_text)
            heading_index += 1
            add_bookmark(paragraph, make_bookmark_name(heading_index), bookmark_id)
            bookmark_id += 1
            i += 1
            continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet_match.group(1))
            i += 1
            continue

        number_match = NUMBER_RE.match(stripped)
        if number_match:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, number_match.group(1))
            i += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped)
        i += 1


def export_markdown_to_docx(src: Path, out: Path) -> None:
    markdown = src.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    render_markdown(doc, markdown)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export markdown to docx with a Word TOC field.")
    parser.add_argument("src", type=Path, help="Source markdown file")
    parser.add_argument("out", type=Path, help="Output docx file")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    export_markdown_to_docx(args.src, args.out)


if __name__ == "__main__":
    main()
