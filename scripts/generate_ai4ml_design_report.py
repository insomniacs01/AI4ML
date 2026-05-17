from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:  # pragma: no cover - diagram fallback is handled at runtime
    Image = None
    ImageDraw = None
    ImageFont = None


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "assets"
TEMPLATE = DOCS_DIR / "template-converted.docx"
OUTPUT_DOCX = DOCS_DIR / "AI4ML-概要设计报告.docx"


REPORT_TITLE = "AI4ML 项目-智算社区平台"
REPORT_SUBTITLE = "概要设计报告"
HEADER_TEXT = "AI4ML 项目-智算社区平台-概要设计报告"
AUTHOR_TEXT = "王文浩，王硕宇，王中颢，李宗欣，姜炳丞，王梓琛"
FINISH_DATE = "2026-05-10"


def remove_all_body_content(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_east_asia_font(run, font_name: str = "宋体", size: Pt | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str, font_name: str = "宋体", size: Pt | None = Pt(10.5), bold: bool | None = None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_east_asia_font(run, font_name=font_name, size=size, bold=bold)


def add_para(document: Document, text: str = "", style: str | None = "Normal", first_line: bool = False):
    paragraph = document.add_paragraph(style=style)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    if text:
        run = paragraph.add_run(text)
        set_east_asia_font(run)
    return paragraph


def add_heading(document: Document, text: str, level: int):
    style = f"Heading {level}"
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    # Keep the template heading style, only set East Asian font to avoid fallback font changes.
    set_east_asia_font(run, font_name="黑体" if level <= 2 else "宋体")
    return paragraph


def add_page_break(document: Document) -> None:
    document.add_page_break()


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_east_asia_font(run, size=Pt(font_size), bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        if widths and i < len(widths):
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths and i < len(widths):
                cells[i].width = Cm(widths[i])
    set_table_borders(table)
    return table


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_east_asia_font(run, size=Pt(10.5), bold=True)


def add_toc(document: Document) -> None:
    title = document.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("目 录")
    set_east_asia_font(run, size=Pt(15), bold=False)

    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr_text)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char)

    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开文档时自动更新。"
    run._r.append(placeholder)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_page_number_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    run._r.append(instr_text)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        run = p.add_run("• " + item)
        set_east_asia_font(run)


def choose_font(size: int, bold: bool = False):
    if ImageFont is None:
        return None
    for font_path in (
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ):
        try:
            return ImageFont.truetype(font_path, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def wrap_text(draw, text: str, font, max_width: int) -> list[str]:
    words = list(text)
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = current + word
        bbox = draw.textbbox((0, 0), candidate, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, text, fill, outline="#2b2b2b", font=None, text_fill="#111111") -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=2)
    lines = wrap_text(draw, text, font, x2 - x1 - 24)
    line_height = (font.size if hasattr(font, "size") else 18) + 5
    total_h = line_height * len(lines)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=text_fill)
        y += line_height


def draw_arrow(draw, start, end, color="#555555", width=3) -> None:
    import math

    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 12
    left = (end[0] - length * math.cos(angle - math.pi / 6), end[1] - length * math.sin(angle - math.pi / 6))
    right = (end[0] - length * math.cos(angle + math.pi / 6), end[1] - length * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=color)


def create_diagrams() -> dict[str, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    if Image is None or ImageDraw is None:
        return {}
    font = choose_font(26)
    small = choose_font(22)
    paths: dict[str, Path] = {}

    architecture = ASSETS_DIR / "ai4ml-design-architecture.png"
    img = Image.new("RGB", (1400, 820), "white")
    draw = ImageDraw.Draw(img)
    boxes = {
        "前端工作台\nReact/Vite": (70, 90, 330, 210),
        "Supabase\nAuth / Postgres / RLS": (550, 90, 860, 210),
        "FastAPI 业务服务\n团队鉴权 / 任务编排 / API": (1030, 90, 1330, 210),
        "AI 连接器\nOpenAI-compatible": (90, 410, 350, 540),
        "MLZero / AutoGluon\n自动建模执行器": (550, 410, 860, 540),
        "本地运行产物\nstorage / logs / output": (1030, 410, 1330, 540),
        "资产与审计\n数据集 / 模型 / 工作流 / 报告": (550, 650, 860, 760),
    }
    fills = ["#f7f7f7", "#eef5ff", "#f7f7f7", "#f9f1e8", "#eef8f0", "#f7f7f7", "#f3eef8"]
    for (text, xy), fill in zip(boxes.items(), fills):
        draw_box(draw, xy, text, fill, font=font)
    draw_arrow(draw, (330, 150), (550, 150))
    draw_arrow(draw, (860, 150), (1030, 150))
    draw_arrow(draw, (1180, 210), (1180, 410))
    draw_arrow(draw, (1030, 475), (860, 475))
    draw_arrow(draw, (550, 475), (350, 475))
    draw_arrow(draw, (705, 540), (705, 650))
    draw_arrow(draw, (1180, 540), (860, 700))
    draw.text((70, 30), "AI4ML 总体架构", font=choose_font(32, True), fill="#111111")
    img.save(architecture)
    paths["architecture"] = architecture

    main_flow = ASSETS_DIR / "ai4ml-design-main-flow.png"
    img = Image.new("RGB", (1500, 620), "white")
    draw = ImageDraw.Draw(img)
    flow = [
        ("创建任务\n自然语言需求", (40, 170, 240, 300), "#f7f7f7"),
        ("上传 CSV\n生成数据画像", (290, 170, 490, 300), "#eef5ff"),
        ("AI 解析\n目标列/类型/指标", (540, 170, 760, 300), "#f9f1e8"),
        ("阶段路由\n选择连接器/模型", (810, 170, 1030, 300), "#f3eef8"),
        ("MLZero 运行\n训练/验证/产物", (1080, 170, 1290, 300), "#eef8f0"),
        ("报告/预测/资产\n复用与审计", (1340, 170, 1480, 300), "#f7f7f7"),
    ]
    for text, xy, fill in flow:
        draw_box(draw, xy, text, fill, font=small)
    for i in range(len(flow) - 1):
        draw_arrow(draw, (flow[i][1][2], 235), (flow[i + 1][1][0], 235))
    draw_box(draw, (520, 420, 980, 540), "失败、缺连接器、缺产物或额度不足时明确失败；不使用伪数据、不静默 fallback", "#fff2cc", font=small)
    draw_arrow(draw, (940, 300), (800, 420), color="#a46a00")
    draw_arrow(draw, (1190, 300), (800, 420), color="#a46a00")
    draw.text((40, 40), "核心建模主流程", font=choose_font(32, True), fill="#111111")
    img.save(main_flow)
    paths["main_flow"] = main_flow

    human_loop = ASSETS_DIR / "ai4ml-design-human-loop.png"
    img = Image.new("RGB", (1300, 620), "white")
    draw = ImageDraw.Draw(img)
    boxes = [
        ("策略触发\nbefore_run / in_run", (60, 180, 300, 300), "#f9f1e8"),
        ("创建复核请求\n分配成员/角色", (380, 180, 620, 300), "#eef5ff"),
        ("人工决策\napprove/revise/block/reject/reassign/skip", (700, 160, 980, 320), "#f3eef8"),
        ("恢复或重跑\n从指定阶段继续", (1060, 180, 1260, 300), "#eef8f0"),
        ("审计与历史\n决策记录/产物引用", (530, 420, 830, 540), "#f7f7f7"),
    ]
    for text, xy, fill in boxes:
        draw_box(draw, xy, text, fill, font=small)
    draw_arrow(draw, (300, 240), (380, 240))
    draw_arrow(draw, (620, 240), (700, 240))
    draw_arrow(draw, (980, 240), (1060, 240))
    draw_arrow(draw, (840, 320), (760, 420))
    draw_arrow(draw, (1080, 300), (830, 470))
    draw.text((60, 40), "人机协同干预流程", font=choose_font(32, True), fill="#111111")
    img.save(human_loop)
    paths["human_loop"] = human_loop

    return paths


def add_image_or_fallback(document: Document, path: Path | None, caption: str, fallback: str) -> None:
    if path and path.exists():
        p = document.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.3))
        add_caption(document, caption)
    else:
        add_para(document, fallback, first_line=False)
        add_caption(document, caption)


def build_report() -> None:
    diagrams = create_diagrams()
    document = Document(str(TEMPLATE))
    remove_all_body_content(document)

    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    header_p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    set_paragraph_text(header_p, HEADER_TEXT, size=Pt(9))
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_footer(document)

    for _ in range(7):
        add_para(document)
    p = document.add_paragraph(style="Normal0")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_TITLE)
    set_east_asia_font(run, size=Pt(16), bold=True)
    p = document.add_paragraph(style="Normal0")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_SUBTITLE)
    set_east_asia_font(run, size=Pt(16), bold=True)
    for _ in range(7):
        add_para(document)

    status = document.add_table(rows=4, cols=3)
    status.style = "Normal Table"
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    status_rows = [
        ["文件状态：\n[√] 草稿\n[  ] 正式发布\n[  ] 正在修改", "文件标识：", "AI4ML-SDD-1"],
        ["文件状态：\n[√] 草稿\n[  ] 正式发布\n[  ] 正在修改", "当前版本：", "V1.0"],
        ["文件状态：\n[√] 草稿\n[  ] 正式发布\n[  ] 正在修改", "作    者：", AUTHOR_TEXT],
        ["文件状态：\n[√] 草稿\n[  ] 正式发布\n[  ] 正在修改", "完成日期：", FINISH_DATE],
    ]
    for r_idx, row in enumerate(status_rows):
        for c_idx, value in enumerate(row):
            set_cell_text(status.rows[r_idx].cells[c_idx], value)
    status.cell(0, 0).merge(status.cell(3, 0))
    set_table_borders(status)

    add_page_break(document)
    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版 本 历 史")
    set_east_asia_font(run, size=Pt(13), bold=True)
    add_table(
        document,
        ["版本/状态", "作者", "审核人", "起止日期", "备注"],
        [
            ["V1.0 / 草稿", AUTHOR_TEXT, "项目组", "2026-05-10", "依据需求规格书、记忆文档、当前代码与数据库结构形成概要设计报告。"],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
        ],
        widths=[2.6, 4.0, 2.5, 3.0, 5.5],
    )

    add_page_break(document)
    add_toc(document)
    add_page_break(document)

    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(REPORT_SUBTITLE)
    set_east_asia_font(run, size=Pt(16), bold=True)

    add_heading(document, "文档介绍", 1)
    add_heading(document, "文档目的", 2)
    add_para(
        document,
        "本文档用于说明 AI4ML“智算社区平台”的概要设计方案，明确系统分层、模块划分、核心流程、接口边界、数据结构和异常处理原则。"
        "文档以《AI4ML 项目需求规格说明书》、当前记忆文档和仓库实现为依据，将需求中定义的团队治理、自然语言建模、AI 工作流、人机协同、资产沉淀与审计能力映射为可实现、可验证的系统设计。",
        first_line=True,
    )
    add_para(
        document,
        "本文档同时约束后续详细设计、编码、测试和验收工作：系统不得使用伪数据冒充真实业务结果，不得在连接器、额度、权限或运行产物缺失时静默兜底；所有正式业务接口除健康检查外均应纳入团队作用域。",
        first_line=True,
    )

    add_heading(document, "文档范围", 2)
    add_para(
        document,
        "本文档覆盖 AI4ML 平台本次交付范围内的前端工作台、FastAPI 业务服务、Supabase 身份与数据层、AI 连接器管理、MLZero/AutoGluon 执行集成、本地运行产物管理、团队治理、配额、审计、资产库、人机协同和代码工作区。"
        "本文档不展开 MLZero 或 AutoGluon 内部算法实现细节，只描述平台如何调用、约束、展示和管理执行底座能力。",
        first_line=True,
    )
    add_para(
        document,
        "首版数据类型以 CSV 表格数据为主，问题类型以分类与回归为主。图像、时序、文本等高级任务可作为后续扩展方向，但不作为本概要设计的强交付边界。",
        first_line=True,
    )

    add_heading(document, "读者对象", 2)
    add_bullets(
        document,
        [
            "项目组前端、后端、AI/Agent 与数据库开发人员。",
            "课程汇报、设计评审、测试验收和部署维护人员。",
            "后续撰写详细设计、测试文档、用户手册和运维说明的成员。",
        ],
    )

    add_heading(document, "参考文献", 2)
    add_table(
        document,
        ["编号", "作者", "文档名称", "单位", "日期"],
        [
            ["[1]", "王文浩，王硕宇，王中颢，李宗欣，姜炳丞，王梓琛", "AI4ML 项目需求规格说明书", "AI4ML 项目组", "2026-04-20"],
            ["[2]", "AI4ML 项目组", "docs/current-memory.md", "AI4ML 项目仓库", "2026-05-09"],
            ["[3]", "AI4ML 项目组", "docs/requirements-coverage-matrix.md", "AI4ML 项目仓库", "2026-05-06"],
            ["[4]", "AI4ML 项目组", "docs/backend-unification-plan.md", "AI4ML 项目仓库", "2026-05-09"],
            ["[5]", "AI4ML 项目组", "README.md", "AI4ML 项目仓库", "2026-05-09"],
            ["[6]", "AI4ML 项目组", "supabase/schema.sql 与 backend/app、frontend/src 当前实现", "AI4ML 项目仓库", "2026-05-10"],
        ],
    )

    add_heading(document, "术语与缩写解释", 2)
    add_table(
        document,
        ["缩写、术语", "解释"],
        [
            ["AI4ML", "AI for Machine Learning，本项目中指面向低代码/零代码建模的团队协作平台。"],
            ["MLZero", "本项目选用的自动建模执行底座，负责生成训练代码、训练模型并输出运行产物。"],
            ["AutoGluon", "自动机器学习框架，当前 MLZero 运行链路中的主要建模能力来源。"],
            ["Supabase", "负责用户认证、团队、关系型数据、RLS 权限和结构化业务数据的服务。"],
            ["FastAPI", "后端业务服务框架，负责团队鉴权后的任务编排、执行调用和产物索引。"],
            ["AI 连接器", "团队可配置的 OpenAI-compatible 模型连接配置，包括 Base URL、模型名、协议和密钥。"],
            ["阶段路由", "按工作流阶段选择 AI 连接器和模型的策略，可由团队默认策略或任务级覆盖提供。"],
            ["Human-in-the-loop", "人机协同机制，允许成员在指定阶段确认、修改、阻断、转交或恢复任务。"],
            ["运行产物", "MLZero/AutoGluon 运行生成的 summary、leaderboard、token_usage、代码、日志、模型文件等。"],
        ],
    )

    add_heading(document, "设计要点", 1)
    add_heading(document, "设计范围", 2)
    add_para(
        document,
        "本概要设计覆盖“团队协作式智能建模工作台”的核心闭环：用户在团队空间内创建任务、上传 CSV、通过 AI 解析任务语义、按阶段路由调用模型连接器、启动 MLZero/AutoGluon 训练，并在运行控制台、模型报告、在线预测、代码工作区、人工复核和资产库中查看真实产物与协作记录。",
        first_line=True,
    )
    add_para(
        document,
        "系统设计以真实数据和真实产物为边界条件。运行成功必须依赖 run_summary、leaderboard、token_usage 等实际产物；若连接器、模型、权限、额度或执行产物缺失，接口和页面必须明确返回失败或“不支持/未接入”。",
        first_line=True,
    )

    add_heading(document, "总体架构", 2)
    add_image_or_fallback(
        document,
        diagrams.get("architecture"),
        "图 2-1 AI4ML 总体架构图",
        "前端工作台 -> Supabase 身份与数据层 -> FastAPI 业务服务 -> MLZero/AutoGluon 执行器 -> 本地运行产物与资产审计。",
    )
    add_para(
        document,
        "前端采用 React 18 + Vite 构建工作台界面；Supabase 负责 Auth、团队成员、RLS、任务、连接器、配额、审计和资产等结构化数据；FastAPI 负责业务 API、团队作用域鉴权、任务生命周期、运行编排、产物索引和文件访问；MLZero/AutoGluon 负责实际建模执行；本地文件系统保存 CSV、运行目录、代码、日志和模型产物。",
        first_line=True,
    )

    add_heading(document, "关键业务功能", 2)
    add_table(
        document,
        ["功能类别", "功能名称、标识符", "描述"],
        [
            ["基础支撑", "FR-BAS-0001 系统健康检查与环境展示", "通过 /api/health 与系统面板展示后端、Provider、执行模式和存储目录状态。"],
            ["基础支撑", "FR-BAS-0002 任务创建、列表与详情查询", "任务按团队隔离，支持创建、查询、详情、语义修正和删除。"],
            ["基础支撑", "FR-BAS-0003 本地执行、Provider 生命周期与失败处理", "后端调用 MLZero executor，记录真实输出目录、失败尝试和错误产物。"],
            ["治理管理", "FR-ADM-0001 用户与权限管理", "基于 Supabase profiles、teams、team_members 与后端 require_team_*_access 实现。"],
            ["治理管理", "FR-ADM-0002 API Token / 资源额度管理", "通过 quota_accounts、token_ledgers 和运行前额度检查控制消耗。"],
            ["治理管理", "FR-ADM-0005 AI 连接器与默认 AI 组合管理", "支持连接器 CRUD、测试、激活、阶段默认路由和任务级覆盖。"],
            ["业务建模", "FR-BIZ-0001 自然语言输入任务需求", "保存任务描述，上传 CSV 后由 AI 解析目标列、任务类型、指标并支持人工修正。"],
            ["业务建模", "FR-BIZ-0002 数据集上传", "校验 CSV 文件并保存 dataset_profile，包括字段类型、缺失值和预览行。"],
            ["业务建模", "FR-BIZ-0003 自动解析需求并启动 AI 工作流", "默认 auto_run 链路先解析语义，再进入 MLZero 运行流程。"],
            ["业务建模", "FR-BIZ-0004 智能体工作进度可视化", "持久化 Agent Runtime、事件、消息、阶段状态和运行观测结果。"],
            ["业务建模", "FR-BIZ-0005 模型训练结果输出", "解析真实 run_summary、leaderboard 与 token_usage，展示最佳模型、指标和候选模型。"],
            ["开发增强", "FR-DEV-0001 阶段级人工复核策略与干预机制", "支持复核请求、决策、转交、驳回、跳过、恢复和阶段重跑提示。"],
        ],
    )

    add_heading(document, "技术实现难点", 2)
    add_bullets(
        document,
        [
            "身份与权限统一：前端身份来自 Supabase，FastAPI 必须信任 Supabase Bearer Token 并按 team_id 校验成员关系。",
            "执行真实性：模型运行必须依赖真实 MLZero/AutoGluon 产物，不能用演示值替代缺失产物。",
            "AI 路由严格性：正式路由只保留主路由字段 stage、connector_id、model_name、config，缺省或无效时直接失败。",
            "长任务观测：运行过程需要从日志、阶段记录、leaderboard、telemetry 和 observer insight 聚合为稳定的前端状态。",
            "人机协同与增量重跑：人工决策需要影响后续运行上下文，并支持从指定阶段重跑或恢复。",
            "文件与结构化数据一致性：Supabase 维护元数据和权限，本地目录保存大文件和运行产物，二者需要通过任务 ID 和输出目录关联。",
        ],
    )

    add_heading(document, "通用设计约定", 2)
    add_bullets(
        document,
        [
            "除 /api/health 外，正式业务接口统一使用 /api/teams/{team_id}/... 路径。",
            "所有任务、连接器、额度、路由、资产、审计和协作记录均必须携带 team_id。",
            "前端页面只展示真实接口数据；未接入或缺失产物时以明确提示展示，不渲染伪状态。",
            "管理员能力由 team_owner 或 admin 角色控制，业务成员与开发成员按页面和 API 权限分层。",
            "CSV 首版使用 UTF-8 文本校验、文件大小校验和空值/字段画像生成。",
            "运行成功以真实 summary、leaderboard、token_usage 等产物是否完整为判定依据。",
        ],
    )

    add_heading(document, "核心建模工作流实现设计", 1)
    add_heading(document, "任务创建与数据上传实现", 2)
    add_para(
        document,
        "任务创建与数据上传是 AI4ML 的入口能力。用户在“开始建模”页面填写任务名称和自然语言描述，可选择阶段级 AI 路由和人工复核策略，并上传单个 CSV 数据集。后端在团队作用域下创建 TaskRecord，CSV 上传后生成 dataset_profile，并根据 auto_run 策略触发语义解析和运行流程。",
        first_line=True,
    )
    add_heading(document, "实现简述及约定", 3)
    add_bullets(
        document,
        [
            "任务创建请求由 TaskCreateRequest 承载，必填 name 和 description，label_column 与 problem_type 可在后续 AI 解析或人工修正后补齐。",
            "上传接口检查文件名、内容类型、大小、空文件、二进制空字节和 UTF-8 解码能力。",
            "数据画像 DatasetProfile 包括行列数、列类型、缺失值数量、缺失比例、样例值和预览行。",
            "任务状态从 draft 进入 uploaded；若自动解析成功并进入运行，则继续流转到 planning、running、completed 或 failed。",
        ],
    )
    add_heading(document, "功能实现流程", 3)
    add_image_or_fallback(
        document,
        diagrams.get("main_flow"),
        "图 3-1 核心建模主流程图",
        "创建任务 -> 上传 CSV -> AI 解析 -> 阶段路由 -> MLZero 运行 -> 报告/预测/资产。",
    )
    add_heading(document, "用户界面设计", 3)
    add_table(
        document,
        ["NO", "类型", "信息内容", "信息表现", "说明"],
        [
            ["1", "必选输入", "任务名称、任务描述", "表单文本框", "用于创建 TaskRecord，长度由后端字段约束控制。"],
            ["2", "必选输入", "CSV 文件", "文件选择控件", "上传后生成 dataset_profile，并作为运行前提。"],
            ["3", "可选输入", "阶段 AI 连接器与模型", "阶段路由表单", "为空时继承团队默认路由；无默认路由则运行失败。"],
            ["4", "可选输入", "人工复核策略", "策略配置表单", "支持 before_run/in_run 触发方式和成员/角色分配。"],
            ["5", "输出", "任务状态、数据画像、下一步动作", "任务详情卡片", "展示真实任务状态和运行建议。"],
        ],
    )
    add_heading(document, "相关接口", 3)
    add_table(
        document,
        ["序号", "类型", "接口编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "内部", "POST /api/teams/{team_id}/tasks", "backend/app/api/routes/task_lifecycle.py", "创建团队作用域任务。"],
            ["2", "内部", "GET /api/teams/{team_id}/tasks", "backend/app/api/routes/task_lifecycle.py", "读取团队任务列表。"],
            ["3", "内部", "GET /api/teams/{team_id}/tasks/{task_id}", "backend/app/api/routes/task_lifecycle.py", "读取任务详情。"],
            ["4", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/dataset", "backend/app/api/routes/task_lifecycle.py", "上传 CSV 并生成数据画像。"],
            ["5", "UI", "TaskForm / TaskCard / App.jsx", "frontend/src", "任务录入、队列展示和详情展示。"],
        ],
    )
    add_heading(document, "出错处理设计", 3)
    add_para(
        document,
        "接口在身份失效、团队无权限、任务不存在、CSV 非法、文件为空、编码错误或运行前提缺失时返回明确错误。前端将错误展示为可读提示，并保留用户已填写的表单上下文。任务运行失败时保留 last_run_attempt、诊断信息和输出目录，供运行控制台和代码工作区继续定位。",
        first_line=True,
    )
    add_heading(document, "出错信息", 3)
    add_table(
        document,
        ["序号", "错误编码", "错误信息", "说明/处理办法"],
        [
            ["1", "TASK-001", "Task not found", "任务不存在或不属于当前团队，提示用户刷新列表或切换团队。"],
            ["2", "DATA-001", "Only CSV upload is supported", "文件类型不符合要求，要求用户重新选择 CSV。"],
            ["3", "DATA-002", "Dataset is empty or not UTF-8", "数据为空或编码不符合首版约束。"],
            ["4", "AUTH-001", "You do not have access to the requested team", "团队鉴权失败，前端提示重新登录或切换团队。"],
            ["5", "RUN-001", "Missing routing connector or model", "阶段路由缺失或不完整，要求配置连接器或团队默认 AI。"],
        ],
    )

    add_heading(document, "AI 解析与 MLZero 运行实现", 2)
    add_heading(document, "实现简述及约定", 3)
    add_para(
        document,
        "AI 解析阶段读取任务描述和数据画像，通过当前阶段路由选择的连接器调用 OpenAI-compatible 模型，生成目标列、任务类型、指标、置信度和结构化需求。运行阶段由 FastAPI 调用 MLZeroExecutor，MLZero/AutoGluon 完成模型搜索、训练和验证，并将结果写入运行目录。",
        first_line=True,
    )
    add_bullets(
        document,
        [
            "正式 AI 路由解析顺序为任务级覆盖优先，再继承团队默认阶段路由。",
            "运行前检查连接器有效性、模型名、额度状态、任务语义和数据集是否齐备。",
            "运行过程记录 workflow_stage_records、task_agent_runs、task_agent_events 和 task_agent_messages。",
            "运行成功后解析 run_summary.json、leaderboard、token_usage.json；缺少关键产物时按失败处理。",
        ],
    )
    add_heading(document, "功能实现流程", 3)
    add_table(
        document,
        ["步骤", "处理对象", "主要动作", "输出"],
        [
            ["1", "TaskStore", "读取任务、数据画像、阶段路由和人工策略", "运行上下文"],
            ["2", "连接器服务", "解析 connector_id、model_name、wire_api 和密钥", "可调用 AI Provider"],
            ["3", "语义分析服务", "生成结构化建模需求并写回 ai_tasks", "structured_requirements / analysis_token_usage"],
            ["4", "MLZeroExecutor", "创建运行目录，调用 AutoGluon/MLZero 训练", "summary、leaderboard、代码、日志"],
            ["5", "产物解析服务", "校验关键产物并生成前端可展示数据", "last_run / task_runs / token_ledgers"],
        ],
    )
    add_heading(document, "用户界面设计", 3)
    add_table(
        document,
        ["NO", "类型", "信息内容", "信息表现", "说明"],
        [
            ["1", "输入", "运行时间限制、重跑阶段", "运行按钮和参数", "支持从指定阶段增量重跑。"],
            ["2", "输出", "AI 解析结果", "任务详情字段", "目标列、任务类型、指标、置信度。"],
            ["3", "输出", "运行状态", "运行控制台", "当前阶段、活动、进度、日志、leaderboard。"],
            ["4", "输出", "失败诊断", "错误提示与日志入口", "展示真实错误产物路径，不伪造成功。"],
        ],
    )
    add_heading(document, "相关接口", 3)
    add_table(
        document,
        ["序号", "类型", "接口编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/analyze", "task_lifecycle.py", "触发语义解析。"],
            ["2", "内部", "PUT /api/teams/{team_id}/tasks/{task_id}/semantic-analysis", "task_lifecycle.py", "人工修正语义并清理旧结果。"],
            ["3", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/run", "task_runtime.py", "启动或重跑 MLZero。"],
            ["4", "内部", "GET /api/teams/{team_id}/tasks/{task_id}/run-progress", "task_runtime.py", "读取真实运行观测。"],
            ["5", "数据", "task_runs / workflow_stage_records / token_ledgers", "supabase/schema.sql", "保存运行、阶段和账本结果。"],
        ],
    )
    add_heading(document, "出错处理设计", 3)
    add_para(
        document,
        "AI 解析或运行失败时，后端保存失败阶段、错误摘要、last_run_attempt、输出目录和错误日志路径。若失败是可恢复的，任务可通过重新运行或从指定阶段重跑继续处理；若触发人工策略，则任务进入 paused_for_review 或 waiting_human。",
        first_line=True,
    )
    add_heading(document, "出错信息", 3)
    add_table(
        document,
        ["序号", "错误编码", "错误信息", "说明/处理办法"],
        [
            ["1", "AI-001", "Connector not found or inactive", "连接器不存在、停用或不属于当前团队。"],
            ["2", "AI-002", "Model name is required", "阶段路由缺少模型名。"],
            ["3", "RUN-002", "MLZero run failed", "运行失败，保留 stdout/stderr 和错误产物。"],
            ["4", "RUN-003", "Missing run_summary or leaderboard", "产物不完整，按失败处理。"],
            ["5", "QUOTA-001", "Token quota exhausted", "额度不足，管理员需调整额度或更换连接器。"],
        ],
    )

    add_heading(document, "运行控制台、报告与代码产物实现", 2)
    add_heading(document, "实现简述及约定", 3)
    add_para(
        document,
        "运行控制台聚合 Agent Runtime、事件流、消息、阶段记录、日志摘要、leaderboard、训练 telemetry 和 observer insight。模型报告页基于真实任务、数据画像、特征重要性文件和运行结果生成解释；代码工作区读取最新运行目录中的可编辑代码、状态、结果、日志和上下文文件。",
        first_line=True,
    )
    add_heading(document, "功能实现流程", 3)
    add_table(
        document,
        ["产物类别", "来源", "展示位置", "设计约束"],
        [
            ["运行进度", "task_run_progress + stage records", "运行控制台", "只展示真实日志和观测状态。"],
            ["模型指标", "run_summary / leaderboard", "任务详情、运行控制台、模型报告", "缺少关键文件则不展示伪结果。"],
            ["报告内容", "TaskModelReportResponse", "模型报告页", "基于真实产物生成解释、局限和路径引用。"],
            ["预测能力", "AutoGluon predictor 或 generated_code.py predict 合约", "在线预测页", "不可加载时返回 supported=false。"],
            ["代码文件", "最新运行目录", "代码工作区", "支持查看、保存版本、下载和单文件重跑。"],
        ],
    )
    add_heading(document, "用户界面设计", 3)
    add_table(
        document,
        ["NO", "界面", "关键输入", "关键输出", "说明"],
        [
            ["1", "运行控制台", "选中任务、刷新操作", "Agent、事件、日志、leaderboard、telemetry", "用于运行态观测。"],
            ["2", "模型报告", "选中任务", "结果摘要、指标说明、特征重要性、风险局限", "用于结果解释。"],
            ["3", "在线预测", "单行 features", "prediction 或 unsupported detail", "只调用真实可加载模型或合约。"],
            ["4", "代码文件", "路径、内容、重跑时间限制", "文件内容、版本历史、stdout/stderr", "开发成员可编辑可写文件。"],
        ],
    )
    add_heading(document, "相关接口", 3)
    add_table(
        document,
        ["序号", "类型", "接口编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "内部", "GET /api/teams/{team_id}/tasks/{task_id}/agent-collaboration", "task_human.py", "读取 Agent 协作快照。"],
            ["2", "内部", "GET /api/teams/{team_id}/tasks/{task_id}/report", "task_artifacts.py", "读取模型报告。"],
            ["3", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/prediction-demo", "task_artifacts.py", "执行在线预测。"],
            ["4", "内部", "GET /api/teams/{team_id}/tasks/{task_id}/code-workspace", "task_artifacts.py", "读取代码工件列表。"],
            ["5", "内部", "PUT /api/teams/{team_id}/tasks/{task_id}/code-workspace/file", "task_artifacts.py", "保存代码文件并记录版本。"],
            ["6", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/code-workspace/rerun", "task_artifacts.py", "重跑可执行 Python 工件。"],
        ],
    )
    add_heading(document, "出错处理设计", 3)
    add_para(
        document,
        "报告、预测和代码工作区均基于最新一次真实运行产物。缺少运行目录、模型文件、预测合约、可编辑文件或权限不足时，接口返回明确错误或 supported=false，前端按页面语义展示“暂不支持/缺少产物”。",
        first_line=True,
    )
    add_heading(document, "出错信息", 3)
    add_table(
        document,
        ["序号", "错误编码", "错误信息", "说明/处理办法"],
        [
            ["1", "ART-001", "Run output directory is missing", "运行目录缺失，提示重新运行任务。"],
            ["2", "ART-002", "Prediction contract is not available", "没有可加载模型或 predict 合约，在线预测不支持。"],
            ["3", "CODE-001", "File is not editable", "只读结果、日志或上下文文件不可保存。"],
            ["4", "CODE-002", "Path is outside run directory", "阻止路径穿越。"],
        ],
    )

    add_heading(document, "治理协作与资产实现设计", 1)
    add_heading(document, "团队权限与成员治理", 2)
    add_heading(document, "实现简述及约定", 3)
    add_para(
        document,
        "系统以 Supabase 为唯一身份和团队真相源。profiles 保存用户资料，teams 保存团队，team_members 保存成员角色和状态。FastAPI 通过 Supabase Bearer Token 识别用户，再结合 team_id 校验成员关系。团队必须至少保留一名 team_owner 或 admin。",
        first_line=True,
    )
    add_heading(document, "功能实现流程", 3)
    add_table(
        document,
        ["步骤", "参与对象", "主要处理", "输出"],
        [
            ["1", "Supabase Auth", "登录/注册后生成 session", "access_token / user_id"],
            ["2", "前端 API 层", "注入 Authorization 与 team_id", "团队作用域请求"],
            ["3", "FastAPI 鉴权依赖", "校验 token 与成员关系", "当前用户与角色上下文"],
            ["4", "团队服务", "处理成员、角色、状态、团队设置和所有权转移", "team_members / audit_logs"],
        ],
    )
    add_heading(document, "用户界面设计", 3)
    add_table(
        document,
        ["NO", "界面", "输入", "输出", "说明"],
        [
            ["1", "登录与团队引导", "邮箱、密码、团队名、邀请码", "会话与 activeTeam", "Supabase 直接处理认证和团队加入。"],
            ["2", "团队与权限页", "角色、成员状态、团队设置", "成员列表、邀请信息、所有权", "管理员或团队所有者执行治理操作。"],
            ["3", "审计日志页", "刷新/筛选", "治理操作记录", "管理员可查看真实 audit_logs。"],
        ],
    )
    add_heading(document, "相关接口", 3)
    add_table(
        document,
        ["序号", "类型", "接口编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "内部", "GET /api/teams/{team_id}/members", "team.py", "成员列表。"],
            ["2", "内部", "PATCH /api/teams/{team_id}/members/{member_id}/role", "team.py", "调整成员角色。"],
            ["3", "内部", "PATCH /api/teams/{team_id}/members/{member_id}/status", "team.py", "冻结、恢复或移除成员。"],
            ["4", "内部", "PATCH /api/teams/{team_id}/settings", "team.py", "更新团队名称、描述和状态。"],
            ["5", "内部", "POST /api/teams/{team_id}/owner/transfer", "team.py", "转移团队所有权。"],
            ["6", "内部", "GET /api/teams/{team_id}/audit-logs", "team.py", "查看审计日志。"],
        ],
    )
    add_heading(document, "出错处理设计", 3)
    add_para(
        document,
        "团队治理接口在非成员访问、非管理员操作、冻结或移除成员访问、转移所有权目标无效、最后一名管理员被移除等情况下拒绝请求。数据库触发器 prevent_last_admin_change 作为后端之外的二次约束。",
        first_line=True,
    )
    add_heading(document, "出错信息", 3)
    add_table(
        document,
        ["序号", "错误编码", "错误信息", "说明/处理办法"],
        [
            ["1", "TEAM-001", "Current user is not a team member", "用户无团队访问权限。"],
            ["2", "TEAM-002", "Admin permission is required", "需要 team_owner 或 admin。"],
            ["3", "TEAM-003", "A team must keep at least one admin", "数据库约束阻止最后管理员变更。"],
            ["4", "TEAM-004", "Target member not found", "操作目标不属于团队。"],
        ],
    )

    add_heading(document, "AI 连接器、阶段路由和配额", 2)
    add_heading(document, "实现简述及约定", 3)
    add_para(
        document,
        "管理员可为团队维护多个 AI 连接器，连接器支持 OpenAI-compatible 协议、chat_completions 或 responses 两种 wire_api。团队默认 AI 组合按阶段保存到 ai_routing_policies，任务可用 stage_routing 覆盖。配额按团队、成员或连接器作用域保存在 quota_accounts，实际消耗写入 token_ledgers。",
        first_line=True,
    )
    add_heading(document, "功能实现流程", 3)
    add_table(
        document,
        ["阶段", "设计处理", "数据表/服务", "说明"],
        [
            ["连接器创建", "保存显示名、Base URL、模型名、协议、加密密钥", "ai_connectors / secret_box", "密钥不明文持久化。"],
            ["连接器测试", "发起最小请求校验可用性", "connector_runtime", "记录 last_test_status。"],
            ["阶段路由", "按阶段保存 connector_id 与 model_name", "ai_routing_policies", "不保留 hidden fallback。"],
            ["运行前额度检查", "按成员/团队/连接器作用域检查剩余额度", "quota_accounts", "不足则阻断运行。"],
            ["用量记账", "Provider usage 或 tokenizer 显式复算", "token_ledgers", "不做固定倍率估算。"],
        ],
    )
    add_heading(document, "用户界面设计", 3)
    add_table(
        document,
        ["NO", "界面", "关键输入", "关键输出", "说明"],
        [
            ["1", "模型连接", "display_name、base_url、model_name、wire_api、api_key", "连接器列表、测试状态、激活状态", "管理员维护。"],
            ["2", "阶段默认 AI", "stage、connector_id、model_name", "团队默认路由", "任务空路由时继承。"],
            ["3", "成员配额", "scope_type、scope_key、token_quota、status、warning_threshold", "额度列表", "支持团队/成员/连接器三类作用域。"],
            ["4", "Token 用量", "刷新", "任务汇总和 ledger 明细", "管理员可查看流水。"],
        ],
    )
    add_heading(document, "相关接口", 3)
    add_table(
        document,
        ["序号", "类型", "接口编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "内部", "GET/POST /api/teams/{team_id}/connectors", "connectors.py", "连接器列表与创建。"],
            ["2", "内部", "PATCH /api/teams/{team_id}/connectors/{connector_id}", "connectors.py", "更新连接器。"],
            ["3", "内部", "POST /api/teams/{team_id}/connectors/{connector_id}/test", "connectors.py", "测试连接器。"],
            ["4", "内部", "POST /api/teams/{team_id}/connectors/{connector_id}/activate", "connectors.py", "激活连接器。"],
            ["5", "内部", "GET/PUT /api/teams/{team_id}/routing", "team.py", "读取/保存阶段默认路由。"],
            ["6", "内部", "GET/POST /api/teams/{team_id}/quotas", "team.py", "额度查询与调整。"],
            ["7", "内部", "GET /api/teams/{team_id}/token-ledgers", "team.py", "用量流水。"],
        ],
    )
    add_heading(document, "出错处理设计", 3)
    add_para(
        document,
        "连接器测试失败、密钥无法解密、路由缺少模型、wire_api 不支持、tokenizer 缺失或额度耗尽时，系统均以明确错误阻断相关动作。激活连接器通过数据库唯一索引保证同一团队只有一个 active 连接器。",
        first_line=True,
    )
    add_heading(document, "出错信息", 3)
    add_table(
        document,
        ["序号", "错误编码", "错误信息", "说明/处理办法"],
        [
            ["1", "CONN-001", "Connector secret key is not configured", "后端未配置加密密钥，无法保存新密钥。"],
            ["2", "CONN-002", "Connector test failed", "连接器不可用，展示 Provider 返回的摘要。"],
            ["3", "ROUTE-001", "Default routing is incomplete", "团队默认路由缺失，要求管理员补齐。"],
            ["4", "QUOTA-002", "Quota account is frozen", "配额被冻结，需管理员恢复。"],
        ],
    )

    add_heading(document, "人机协同复核机制", 2)
    add_heading(document, "实现简述及约定", 3)
    add_para(
        document,
        "人机协同用于在建模流程中引入可控干预。任务可保存 interaction_policies，策略命中后创建 human_interaction_requests 并暂停任务。成员可在复核待办中 approve、revise、block、reject、reassign 或 skip，决策会进入任务上下文和审计记录，必要时引导从指定阶段重跑。",
        first_line=True,
    )
    add_heading(document, "功能实现流程", 3)
    add_image_or_fallback(
        document,
        diagrams.get("human_loop"),
        "图 4-1 人机协同干预流程图",
        "策略触发 -> 创建复核请求 -> 人工决策 -> 恢复或重跑 -> 审计与历史。",
    )
    add_heading(document, "用户界面设计", 3)
    add_table(
        document,
        ["NO", "界面", "关键输入", "关键输出", "说明"],
        [
            ["1", "任务表单", "interaction_policies", "任务级复核策略", "创建任务时保存。"],
            ["2", "复核待办", "决策动作、说明、产物路径、转交对象", "请求状态、决策历史、恢复入口", "只处理人工决策。"],
            ["3", "运行控制台", "刷新", "人工节点事件", "只展示运行态，不承载决策表单。"],
        ],
    )
    add_heading(document, "相关接口", 3)
    add_table(
        document,
        ["序号", "类型", "接口编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "内部", "GET /api/teams/{team_id}/tasks/{task_id}/human-collaboration", "task_human.py", "读取复核请求与决策历史。"],
            ["2", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/human-requests", "task_human.py", "创建人工复核请求。"],
            ["3", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/human-requests/{request_id}/decision", "task_human.py", "提交人工决策。"],
            ["4", "内部", "POST /api/teams/{team_id}/tasks/{task_id}/resume", "task_human.py", "恢复等待中的任务。"],
        ],
    )
    add_heading(document, "出错处理设计", 3)
    add_para(
        document,
        "复核请求状态必须防止重复决策；已关闭、过期或非当前团队的请求不允许提交。revise/reject 等动作可标记 rerun_from_stage，恢复时由运行服务读取下一轮人类指导上下文。",
        first_line=True,
    )
    add_heading(document, "出错信息", 3)
    add_table(
        document,
        ["序号", "错误编码", "错误信息", "说明/处理办法"],
        [
            ["1", "HITL-001", "Human request not found", "请求不存在或不属于当前任务。"],
            ["2", "HITL-002", "Request is already closed", "请求已处理，不允许重复决策。"],
            ["3", "HITL-003", "Resume is not allowed", "任务当前状态不允许恢复。"],
            ["4", "HITL-004", "Reassign target is invalid", "转交成员或角色无效。"],
        ],
    )

    add_heading(document, "资产库与复用机制", 2)
    add_heading(document, "实现简述及约定", 3)
    add_para(
        document,
        "资产库统一登记 dataset、model、workflow、report 四类资产，记录标题、描述、存储路径、分类、标签、可见性、版本、来源任务、模型卡和元数据。管理员可审核和发布资产，成员可 Fork 资产形成副本。当前版本定位为团队内资产登记与审核台账，而非完整文件托管系统或跨团队公开市场。",
        first_line=True,
    )
    add_heading(document, "功能实现流程", 3)
    add_table(
        document,
        ["步骤", "处理对象", "主要动作", "输出"],
        [
            ["1", "任务产物", "从当前任务沉淀数据集、模型、工作流或报告资产", "platform_assets 记录"],
            ["2", "资产审核", "管理员调整 review_status、category、tags、visibility", "审计日志"],
            ["3", "资产发布", "设置 visibility 与 published_at", "可见资产"],
            ["4", "资产 Fork", "复制资产元数据并记录 source_asset_id", "独立副本"],
        ],
    )
    add_heading(document, "用户界面设计", 3)
    add_table(
        document,
        ["NO", "界面", "输入", "输出", "说明"],
        [
            ["1", "资产中心", "资产类型、标题、描述、路径、标签", "资产列表", "团队成员登记资产。"],
            ["2", "资产审核", "review_status、category、tags", "状态变更", "管理员维护。"],
            ["3", "发布/Fork", "visibility、标题、描述", "发布资产或副本", "用于社区沉淀和复用。"],
        ],
    )
    add_heading(document, "相关接口", 3)
    add_table(
        document,
        ["序号", "类型", "接口编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "内部", "GET /api/teams/{team_id}/assets", "team.py", "资产列表。"],
            ["2", "内部", "POST /api/teams/{team_id}/assets", "team.py", "创建资产。"],
            ["3", "内部", "POST /api/teams/{team_id}/assets/{asset_id}/review", "team.py", "审核资产。"],
            ["4", "内部", "POST /api/teams/{team_id}/assets/{asset_id}/publish", "team.py", "发布资产。"],
            ["5", "内部", "POST /api/teams/{team_id}/assets/{asset_id}/fork", "team.py", "Fork 资产副本。"],
        ],
    )
    add_heading(document, "出错处理设计", 3)
    add_para(
        document,
        "资产操作检查团队权限、资产归属、资产类型和可见性枚举。Fork 后副本必须拥有独立 ID 和当前团队归属，保留 source_asset_id 作为来源追溯。发布与审核操作写入 audit_logs。",
        first_line=True,
    )
    add_heading(document, "出错信息", 3)
    add_table(
        document,
        ["序号", "错误编码", "错误信息", "说明/处理办法"],
        [
            ["1", "ASSET-001", "Asset not found", "资产不存在或不可见。"],
            ["2", "ASSET-002", "Unsupported asset type", "资产类型不在 dataset/model/workflow/report 中。"],
            ["3", "ASSET-003", "Review permission required", "审核或发布需要管理员权限。"],
        ],
    )

    add_heading(document, "接口设计", 1)
    add_para(
        document,
        "接口采用 REST 风格，默认 JSON 请求/响应。前端调用 FastAPI 时携带 Supabase access_token，后端通过 team_id 路径参数限定业务作用域。健康检查接口不要求团队上下文；其他业务接口统一挂载到 /api/teams/{team_id}/...。",
        first_line=True,
    )
    add_heading(document, "接口总体约定", 2)
    add_bullets(
        document,
        [
            "认证方式：Authorization: Bearer <supabase_access_token>。",
            "团队作用域：路径参数 team_id 必须与当前用户所在团队匹配。",
            "错误返回：使用 HTTP 状态码表达权限、参数、资源不存在、运行失败等类别，并在 detail 中提供可读说明。",
            "数据格式：请求体和响应体使用 Pydantic 模型约束；文件上传使用 multipart/form-data。",
            "安全约束：禁止路径穿越，连接器密钥加密保存，团队数据由 Supabase RLS 与后端鉴权双重限制。",
        ],
    )
    add_heading(document, "核心任务接口", 2)
    add_table(
        document,
        ["接口编号/名称", "接口方式", "用途", "输入", "输出"],
        [
            ["TASK-LIST", "GET /api/teams/{team_id}/tasks", "获取任务列表", "team_id", "TaskListResponse"],
            ["TASK-CREATE", "POST /api/teams/{team_id}/tasks", "创建任务", "TaskCreateRequest", "TaskRecord"],
            ["TASK-DETAIL", "GET /api/teams/{team_id}/tasks/{task_id}", "获取任务详情", "task_id", "TaskRecord"],
            ["TASK-CONFIG", "PUT /api/teams/{team_id}/tasks/{task_id}/workflow-config", "更新阶段路由和人工策略", "TaskWorkflowConfigUpdateRequest", "TaskRecord"],
            ["TASK-SEMANTIC", "PUT /api/teams/{team_id}/tasks/{task_id}/semantic-analysis", "人工修正语义", "TaskSemanticUpdateRequest", "TaskRecord"],
            ["TASK-DATASET", "POST /api/teams/{team_id}/tasks/{task_id}/dataset", "上传 CSV", "UploadFile", "TaskRecord"],
            ["TASK-ANALYZE", "POST /api/teams/{team_id}/tasks/{task_id}/analyze", "AI 语义解析", "task_id", "TaskRecord"],
            ["TASK-RUN", "POST /api/teams/{team_id}/tasks/{task_id}/run", "启动或重跑 MLZero", "TaskRunRequest", "TaskRecord"],
        ],
    )
    add_heading(document, "运行与产物接口", 2)
    add_table(
        document,
        ["接口编号/名称", "接口方式", "用途", "输入", "输出"],
        [
            ["RUN-PROGRESS", "GET /api/teams/{team_id}/tasks/{task_id}/run-progress", "读取运行进度", "task_id", "TaskRunProgressResponse"],
            ["AGENT-COLLAB", "GET /api/teams/{team_id}/tasks/{task_id}/agent-collaboration", "读取 Agent 协作", "task_id", "TaskAgentCollaborationResponse"],
            ["HUMAN-COLLAB", "GET /api/teams/{team_id}/tasks/{task_id}/human-collaboration", "读取人工复核", "task_id", "TaskHumanCollaborationResponse"],
            ["REPORT", "GET /api/teams/{team_id}/tasks/{task_id}/report", "生成模型报告", "task_id", "TaskModelReportResponse"],
            ["PREDICT", "POST /api/teams/{team_id}/tasks/{task_id}/prediction-demo", "在线预测", "TaskPredictionDemoRequest", "TaskPredictionDemoResponse"],
            ["CODE-LIST", "GET /api/teams/{team_id}/tasks/{task_id}/code-workspace", "读取代码工件", "task_id", "TaskCodeWorkspaceResponse"],
            ["CODE-FILE", "GET/PUT /api/teams/{team_id}/tasks/{task_id}/code-workspace/file", "读取或保存文件", "path/content", "TaskCodeArtifactContentResponse"],
            ["CODE-RERUN", "POST /api/teams/{team_id}/tasks/{task_id}/code-workspace/rerun", "重跑 Python 工件", "TaskCodeArtifactRerunRequest", "TaskCodeArtifactRerunResponse"],
        ],
    )
    add_heading(document, "团队治理接口", 2)
    add_table(
        document,
        ["接口编号/名称", "接口方式", "用途", "输入", "输出"],
        [
            ["MEMBERS", "GET /api/teams/{team_id}/members", "读取团队成员", "team_id", "TeamMembersResponse"],
            ["SETTINGS", "GET/PATCH /api/teams/{team_id}/settings", "读取或更新团队设置", "TeamSettingsUpdateRequest", "TeamSettingsResponse"],
            ["OWNER-TRANSFER", "POST /api/teams/{team_id}/owner/transfer", "转移所有权", "TeamOwnershipTransferRequest", "TeamOwnershipTransferResponse"],
            ["INVITE", "POST /api/teams/{team_id}/members/invite", "生成邀请信息", "TeamInviteRequest", "TeamInviteResponse"],
            ["ROLE", "PATCH /api/teams/{team_id}/members/{member_id}/role", "更新成员角色", "TeamMemberRoleUpdateRequest", "TeamMemberRoleUpdateResponse"],
            ["STATUS", "PATCH /api/teams/{team_id}/members/{member_id}/status", "更新成员状态", "TeamMemberStatusUpdateRequest", "TeamMemberStatusUpdateResponse"],
            ["AUDIT", "GET /api/teams/{team_id}/audit-logs", "读取审计日志", "team_id", "AuditLogsResponse"],
        ],
    )
    add_heading(document, "连接器、路由、配额与资产接口", 2)
    add_table(
        document,
        ["接口编号/名称", "接口方式", "用途", "输入", "输出"],
        [
            ["CONNECTORS", "GET/POST /api/teams/{team_id}/connectors", "连接器列表和创建", "ConnectorCreateRequest", "ConnectorListResponse/ConnectorRecord"],
            ["CONNECTOR-TEST", "POST /api/teams/{team_id}/connectors/{connector_id}/test", "测试连接器", "connector_id", "ConnectorTestResponse"],
            ["CONNECTOR-ACTIVE", "POST /api/teams/{team_id}/connectors/{connector_id}/activate", "激活连接器", "connector_id", "ConnectorActivateResponse"],
            ["ROUTING", "GET/PUT /api/teams/{team_id}/routing", "阶段默认 AI 路由", "AIRoutingPoliciesUpdateRequest", "AIRoutingPoliciesResponse"],
            ["QUOTAS", "GET/POST /api/teams/{team_id}/quotas", "额度查询和调整", "TeamQuotaScopeAdjustRequest", "TeamQuotasResponse"],
            ["USAGE", "GET /api/teams/{team_id}/usage", "团队用量汇总", "team_id", "TeamTokenUsageResponse"],
            ["LEDGERS", "GET /api/teams/{team_id}/token-ledgers", "用量流水", "team_id", "TokenLedgersResponse"],
            ["ASSETS", "GET/POST /api/teams/{team_id}/assets", "资产列表和创建", "PlatformAssetCreateRequest", "PlatformAssetsResponse"],
        ],
    )

    add_heading(document, "数据结构设计", 1)
    add_para(
        document,
        "系统数据结构分为 Supabase 结构化业务数据和本地文件产物两部分。Supabase 保存身份、团队、任务、连接器、路由、运行摘要、账本、阶段、协作、复核、资产和审计；本地文件系统保存 CSV、MLZero 输出目录、模型、代码、日志和可下载工件。",
        first_line=True,
    )
    add_heading(document, "核心数据对象", 2)
    add_table(
        document,
        ["序号", "类型", "数据编号/名称", "来源/文档编号", "说明"],
        [
            ["1", "数据库表", "profiles", "supabase/schema.sql", "用户资料，与 Supabase auth.users 关联。"],
            ["2", "数据库表", "teams", "supabase/schema.sql", "团队基本信息、邀请码、状态和创建人。"],
            ["3", "数据库表", "team_members", "supabase/schema.sql", "团队成员、角色、状态和邀请关系。"],
            ["4", "数据库表", "ai_connectors", "supabase/schema.sql", "AI 连接器配置和测试状态。"],
            ["5", "数据库表", "ai_tasks", "supabase/schema.sql", "任务主表，聚合任务语义、数据集、路由、交互策略和最近运行结果。"],
            ["6", "数据库表", "task_runs", "supabase/schema.sql", "任务运行记录，保存输出目录、最佳模型、指标、leaderboard 和 token_usage。"],
            ["7", "数据库表", "workflow_stage_records", "supabase/schema.sql", "阶段状态、模型选择、耗时、日志摘要和产物引用。"],
            ["8", "数据库表", "task_agent_runs / events / messages", "supabase/schema.sql", "Agent Runtime、事件流和协作消息。"],
            ["9", "数据库表", "human_interaction_requests", "supabase/schema.sql", "人工复核请求、分配、状态、决策和超时。"],
            ["10", "数据库表", "quota_accounts / token_ledgers", "supabase/schema.sql", "额度账户和 token 消耗流水。"],
            ["11", "数据库表", "platform_assets", "supabase/schema.sql", "数据集、模型、工作流和报告资产。"],
            ["12", "数据库表", "audit_logs", "supabase/schema.sql", "团队治理、资产、复核等关键操作审计。"],
        ],
    )
    add_heading(document, "ai_tasks 数据结构", 2)
    add_table(
        document,
        ["序号", "字段", "长度", "类型", "说明"],
        [
            ["1", "id", "8+", "text", "任务唯一标识，系统生成。"],
            ["2", "team_id", "uuid", "uuid", "所属团队，外键 teams.id。"],
            ["3", "created_by / creator_user_id", "uuid", "uuid", "创建人和业务创建人。"],
            ["4", "name", "1-80", "text", "任务名称。"],
            ["5", "description", "1-500", "text", "自然语言任务描述。"],
            ["6", "label_column", "120", "text", "目标列，可由 AI 解析或人工修正。"],
            ["7", "problem_type", "enum", "text", "classification 或 regression。"],
            ["8", "status", "enum", "text", "draft、uploaded、planning、running、paused_for_review、waiting_human、completed、failed、published。"],
            ["9", "dataset_filename / dataset_path", "文本", "text", "上传数据集名称和路径。"],
            ["10", "dataset_profile", "json", "jsonb", "数据画像。"],
            ["11", "structured_requirements", "json", "jsonb", "AI 解析或人工修正后的结构化需求。"],
            ["12", "stage_routing", "json", "jsonb", "任务级阶段路由覆盖。"],
            ["13", "interaction_policies", "json", "jsonb", "任务级人工复核策略。"],
            ["14", "last_run / last_run_attempt", "json", "jsonb", "最近成功运行或失败尝试摘要。"],
        ],
    )
    add_heading(document, "运行与阶段数据结构", 2)
    add_table(
        document,
        ["序号", "数据编号/名称", "关键字段", "说明"],
        [
            ["1", "task_runs", "team_id、task_id、status、output_dir、best_model、metric_name、metric_value、leaderboard、token_usage", "记录一次运行的摘要和核心产物索引。"],
            ["2", "workflow_stage_records", "stage、status、selected_connector_id、model_name、summary、artifact_refs、started_at、finished_at、duration_seconds", "记录每个工作流阶段的执行结果。"],
            ["3", "task_agent_runs", "agent_id、stage、name、role、status、progress、current_task、worker_id", "记录当前 Agent Runtime 快照。"],
            ["4", "task_agent_events", "agent_id、stage、kind、status、text、artifact_refs、created_at", "记录事件流。"],
            ["5", "task_agent_messages", "from_agent_id、to_agent_id、message_type、status、content、correlation_id", "记录 Agent 间协作与人工节点消息。"],
        ],
    )
    add_heading(document, "治理与协作数据结构", 2)
    add_table(
        document,
        ["序号", "数据编号/名称", "关键字段", "说明"],
        [
            ["1", "team_members", "team_id、user_id、role、member_status、invited_by", "角色和成员状态。"],
            ["2", "ai_connectors", "display_name、base_url、model_name、wire_api、api_key、is_active、last_test_status", "AI Provider 连接信息。"],
            ["3", "ai_routing_policies", "team_id、stage、connector_id、model_name、config", "团队默认阶段路由。"],
            ["4", "quota_accounts", "scope_type、scope_key、token_quota、token_used、status、warning_threshold", "额度账户。"],
            ["5", "token_ledgers", "phase、stage_key、source_key、model_name、input_tokens、output_tokens、total_tokens", "消耗流水。"],
            ["6", "human_interaction_requests", "stage、status、assigned_to、assignee_type、payload、decision、timeout_at", "人工复核请求与决策。"],
            ["7", "platform_assets", "asset_type、title、storage_path、category、tags、visibility、source_task_id、source_asset_id、model_card", "资产登记、发布和 Fork。"],
            ["8", "audit_logs", "actor_id、action、resource_type、resource_id、detail、created_at", "关键操作审计。"],
        ],
    )
    add_heading(document, "文件产物存储结构", 2)
    add_para(
        document,
        "本地文件系统用于保存上传 CSV 与 MLZero 运行产物。当前任务数据集和执行结果按任务与运行目录关联，Windows 默认运行产物目录可位于 %LOCALAPPDATA%\\AI4ML\\mlzero_runs，仓库内也包含 storage、logs、output 等辅助目录。文件路径只作为后端内部索引和受控下载入口，不直接暴露任意文件系统访问能力。",
        first_line=True,
    )
    add_table(
        document,
        ["序号", "路径/产物", "用途", "约束"],
        [
            ["1", "storage/tasks/<task_id>/dataset.csv", "任务上传数据集", "仅允许 CSV，元数据按 team_id 存入 Supabase。"],
            ["2", "mlzero_runs/<run_id>/run_summary.json", "最佳模型与指标摘要", "缺失时运行不判定成功。"],
            ["3", "mlzero_runs/<run_id>/leaderboard.*", "候选模型结果", "用于 leaderboard 和报告。"],
            ["4", "mlzero_runs/<run_id>/token_usage.json", "模型调用消耗", "用于 token_ledgers。"],
            ["5", "mlzero_runs/<run_id>/generated_code.py", "生成代码或预测合约", "代码工作区可读；可编辑文件保存版本。"],
            ["6", "mlzero_runs/<run_id>/stdout/stderr/logs", "运行日志与错误诊断", "运行控制台和失败处理读取。"],
        ],
    )
    add_heading(document, "数据规则与约束", 2)
    add_bullets(
        document,
        [
            "所有团队级业务表启用 RLS，并通过 is_member_of_team 或 is_team_admin 限制访问。",
            "team_members 通过 prevent_last_admin_change 触发器保证团队至少保留一名管理员或所有者。",
            "ai_connectors 通过唯一索引保证同一团队最多一个 active 连接器。",
            "ai_routing_policies 以 (team_id, stage) 唯一约束保存默认阶段路由。",
            "token_ledgers 以 (team_id, task_id, phase, source_key) 避免同一来源重复记账。",
            "platform_assets 通过 source_task_id 和 source_asset_id 支持任务沉淀与 Fork 追溯。",
            "文件访问必须限定在任务运行目录内，避免路径穿越和跨任务读取。",
        ],
    )

    document.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_report()
    print(OUTPUT_DOCX)
