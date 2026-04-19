from __future__ import annotations

import argparse
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HEADING_RE = re.compile(r"^(#+)\s+(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+\s*$")
BULLET_RE = re.compile(r"^-\s+(.*)$")
NUMBER_RE = re.compile(r"^\d+\.\s+(.*)$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export SRS markdown to docx with a sample-like layout.")
    parser.add_argument("src", type=Path, help="Source markdown file")
    parser.add_argument("out", type=Path, help="Output docx file")
    parser.add_argument("--project-name", default="“智算”AI4ML 社区", help="Project display name")
    parser.add_argument("--doc-id", default="AI4ML-FR-1", help="Document id")
    parser.add_argument("--version", default="1.0", help="Document version")
    parser.add_argument("--author", default="王中颢", help="Document author")
    parser.add_argument("--date", default="2026-04-06", help="Completion date")
    return parser.parse_args()


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Songti SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document) -> None:
    styles = doc.styles

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = styles[style_name]
        style.font.name = "Songti SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")

    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 4"].font.size = Pt(11)

    for toc_style_name in ["TOC 1", "TOC 2", "TOC 3"]:
        if toc_style_name in styles:
            style = styles[toc_style_name]
            style.font.name = "Songti SC"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
            if toc_style_name == "TOC 1":
                style.font.size = Pt(11.5)
            elif toc_style_name == "TOC 2":
                style.font.size = Pt(11)
            else:
                style.font.size = Pt(10.5)

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "Songti SC")
    fonts.set(qn("w:ascii"), "Songti SC")
    fonts.set(qn("w:hAnsi"), "Songti SC")
    rpr.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rpr.append(underline)

    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def collect_headings(markdown: str) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        stripped = line.strip()
        match = HEADING_RE.match(stripped)
        if not match:
            continue
        level = len(match.group(1))
        text = match.group(2).strip()
        if level == 1:
            continue
        if level > 4:
            continue
        mapped_level = max(1, level - 1)
        headings.append((mapped_level, text))
    return headings


def estimate_toc_pages(num_entries: int, entries_per_page: int = 34) -> int:
    return max(1, math.ceil((num_entries + 1) / entries_per_page))


def estimate_heading_pages(markdown: str, start_page: int, lines_per_page: float = 39.0) -> list[int]:
    lines = markdown.splitlines()
    page = start_page
    used = 0.0
    in_code = False
    pages: list[int] = []
    i = 0

    def block_lines(text: str, chars_per_line: int = 34, extra: float = 0.4) -> float:
        if not text.strip():
            return 0.0
        return max(1.0, math.ceil(len(text) / chars_per_line) + extra)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            height = 1.0
            if used + height > lines_per_page:
                page += 1
                used = 0.0
            used += height
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1].strip()):
            row_count = 1
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_count += 1
                i += 1
            height = row_count * 1.25 + 0.5
            if used + height > lines_per_page:
                page += 1
                used = 0.0
            used += height
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level > 1 and level <= 4:
                height = block_lines(text, chars_per_line=26, extra=0.8)
                if used + height > lines_per_page:
                    page += 1
                    used = 0.0
                pages.append(page)
                used += height
            elif level == 5:
                height = block_lines(text, chars_per_line=28, extra=0.5)
                if used + height > lines_per_page:
                    page += 1
                    used = 0.0
                used += height
            i += 1
            continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            height = block_lines(bullet_match.group(1), chars_per_line=32, extra=0.4)
            if used + height > lines_per_page:
                page += 1
                used = 0.0
            used += height
            i += 1
            continue

        number_match = NUMBER_RE.match(stripped)
        if number_match:
            height = block_lines(number_match.group(1), chars_per_line=32, extra=0.4)
            if used + height > lines_per_page:
                page += 1
                used = 0.0
            used += height
            i += 1
            continue

        height = block_lines(stripped, chars_per_line=34, extra=0.4)
        if used + height > lines_per_page:
            page += 1
            used = 0.0
        used += height
        i += 1

    return pages


def add_cover_page(doc: Document, args: argparse.Namespace) -> None:
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(args.project_name)
    set_run_font(run, size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"({args.doc_id})")
    set_run_font(run, size=16, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("需求规格说明书")
    set_run_font(run, size=18, bold=True)

    for _ in range(7):
        doc.add_paragraph()

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    rows = [
        ("文件状态：\n[  ] 草稿\n[√] 正式发布\n[  ] 正在修改", "文件标识：", args.doc_id),
        ("", "当前版本：", args.version),
        ("", "作    者：", args.author),
        ("", "完成日期：", args.date),
    ]

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for idx, line in enumerate(text.split("\n")):
                if idx:
                    p.add_run("\n")
                run = p.add_run(line)
                set_run_font(run, size=10.5)

    doc.add_page_break()


def add_version_history_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版 本 历 史")
    set_run_font(run, size=14, bold=True)

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    values = [
        ["版本 / 状态", "作者", "更新日期", "更新说明"],
        ["0.1 / 草稿", "王中颢", "2026-04-01", "完成项目需求梳理、底座分析与初版需求文档整理。"],
        ["1.0 / 正式发布", "王中颢", "2026-04-06", "按课程 Week 6 提交要求形成完整需求规格说明书，并统一为正式提交版样式。"],
    ]

    for r_idx, row in enumerate(values):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, size=10.5, bold=(r_idx == 0))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("字体颜色说明")
    set_run_font(run, size=12, bold=True)

    p = doc.add_paragraph()
    run = p.add_run("黑色字体内容表示本版本确认后的正式需求内容。")
    set_run_font(run, size=10.5)
    p = doc.add_paragraph()
    run = p.add_run("红色字体内容保留给后续修订版新增内容，本版本暂不使用红色修订标记。")
    set_run_font(run, size=10.5)


def add_toc_page(doc: Document, headings: list[tuple[int, str]], heading_pages: list[int]) -> None:
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目 录")
    set_run_font(run, size=14, bold=True)

    for index, (level, text) in enumerate(headings, start=1):
        style_name = f"TOC {min(level, 3)}" if f"TOC {min(level, 3)}" in doc.styles else "Normal"
        p = doc.add_paragraph(style=style_name)
        p.paragraph_format.left_indent = Inches(0.22 * (level - 1))
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.1),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.SPACES,
        )
        add_internal_hyperlink(p, text, f"toc_anchor_{index}")
        run = p.add_run("\t" + str(heading_pages[index - 1]))
        set_run_font(run, size=10.5)


def add_body_title(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("软件需求规格书")
    set_run_font(run, size=14, bold=True)


def add_inline(paragraph, text: str) -> None:
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j != -1:
                run = paragraph.add_run(text[i + 2 : j])
                set_run_font(run, bold=True)
                i = j + 2
                continue
        if text.startswith("`", i):
            j = text.find("`", i + 1)
            if j != -1:
                run = paragraph.add_run(text[i + 1 : j])
                run.font.name = "Menlo"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
                run.font.size = Pt(10.5)
                i = j + 1
                continue
        next_pos = len(text)
        for token in ("**", "`"):
            pos = text.find(token, i)
            if pos != -1:
                next_pos = min(next_pos, pos)
        run = paragraph.add_run(text[i:next_pos])
        set_run_font(run)
        i = next_pos


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(10.5)
                if r_idx == 0:
                    run.bold = True


def render_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    in_code = False
    code_buffer: list[str] = []
    title_skipped = False
    bookmark_id = 1
    toc_heading_index = 0
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                for idx, code_line in enumerate(code_buffer):
                    if idx:
                        p.add_run("\n")
                    run = p.add_run(code_line)
                    run.font.name = "Menlo"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
                    run.font.size = Pt(10)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[i + 1].strip()):
            table_lines = [line]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in table_lines]
            add_table(doc, rows)
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1 and not title_skipped:
                title_skipped = True
                i += 1
                continue

            if level == 2:
                style_name = "Heading 1"
            elif level == 3:
                style_name = "Heading 2"
            elif level == 4:
                style_name = "Heading 3"
            else:
                style_name = "Heading 4"

            p = doc.add_paragraph(style=style_name)
            add_inline(p, text)

            if 2 <= level <= 4:
                toc_heading_index += 1
                add_bookmark(p, f"toc_anchor_{toc_heading_index}", bookmark_id)
                bookmark_id += 1

            i += 1
            continue

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet_match.group(1))
            i += 1
            continue

        number_match = NUMBER_RE.match(stripped)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, number_match.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


def main() -> None:
    args = parse_args()
    markdown = args.src.read_text(encoding="utf-8")
    headings = collect_headings(markdown)
    toc_pages = estimate_toc_pages(len(headings))
    body_start_page = 3 + toc_pages
    heading_pages = estimate_heading_pages(markdown, start_page=body_start_page)

    doc = Document()
    configure_document(doc)
    add_cover_page(doc, args)
    add_version_history_page(doc)
    add_toc_page(doc, headings, heading_pages)
    add_body_title(doc)
    render_body(doc, markdown)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(args.out)


if __name__ == "__main__":
    main()
