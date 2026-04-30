# -*- coding: utf-8 -*-
from docx import Document
import comtypes.client


DOCX_PATH = r"D:\333\AI4ML\周报\MON-智算社区-2-周报-0427.docx"
DOC_PATH = r"D:\333\AI4ML\周报\MON-智算社区-2-周报-0427.doc"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
        replace_paragraph_text(paragraph, text)
        return

    replace_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        replace_paragraph_text(paragraph, "")


def renumber_milestone_table() -> None:
    document = Document(DOCX_PATH)
    table = document.tables[2]

    for row_index in range(1, len(table.rows)):
        replace_cell_text(table.cell(row_index, 0), str(row_index))

    document.save(DOCX_PATH)


def export_doc() -> None:
    word = comtypes.client.CreateObject("Word.Application", dynamic=True)
    word.Visible = False
    word.DisplayAlerts = 0
    opened = None
    try:
        opened = word.Documents.Open(DOCX_PATH)
        opened.SaveAs(DOC_PATH, 0)
    finally:
        if opened is not None:
            opened.Close(False)
        word.Quit()


def main() -> None:
    renumber_milestone_table()
    export_doc()

    check = Document(DOCX_PATH)
    table = check.tables[2]
    print([table.cell(1, col).text for col in range(5)])
    print([table.cell(2, col).text for col in range(5)])
    print([table.cell(len(table.rows) - 1, col).text for col in range(5)])


if __name__ == "__main__":
    main()
